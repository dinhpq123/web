#!/usr/bin/env python3
"""
Quawaco IOC Center – System Specification Generator
Generates docs/Quawaco_IOC_Spec.docx
"""
import os, sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")

# ── Paths ──────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
IMGS = SCRIPT_DIR / "screenshots"
OUT  = SCRIPT_DIR / "Quawaco_IOC_Spec.docx"

def add_image(doc, name, width_inches=6.0):
    p = IMGS / name
    if not p.exists():
        matches = list(IMGS.glob(f"{p.stem}_*.png"))
        if matches:
            p = sorted(matches)[-1]
    if p.exists():
        doc.add_picture(str(p), width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print(f"WARNING: Image not found: {p}")

def create_doc():
    doc = Document()

    # ── Cover Page ──────────────────────────────────────────────────────────
    doc.add_heading('QUAWACO IOC CENTER', 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Đặc tả Hệ thống Điều hành Thông minh v1.0')
    run.font.size = Pt(20)
    run.bold = True
    
    p = doc.add_paragraph(f'Ngày lập: {datetime.now().strftime("%d/%m/%Y")}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # ── Introduction ───────────────────────────────────────────────────────
    doc.add_heading('1. Giới thiệu tổng quan', level=1)
    doc.add_paragraph(
        'Hệ thống IOC Center (Intelligent Operations Center) của Quawaco là nền tảng quản trị số toàn diện, '
        'tích hợp dữ liệu từ nhiều nguồn khác nhau bao gồm SCADA, GIS, Camera và CRM để cung cấp cái nhìn 360 độ '
        'về hoạt động vận hành cấp nước của Công ty Cổ phần Nước sạch Quảng Ninh.'
    )
    
    add_image(doc, '01_login_page.png')
    doc.add_paragraph('Màn hình đăng nhập tích hợp xác thực đa yếu tố (2FA) và phân quyền 7 vai trò (RBAC).')

    # ── Command Center ─────────────────────────────────────────────────────
    doc.add_heading('2. Phân hệ Điều hành Tập trung (Command Center)', level=1)
    doc.add_paragraph(
        'Đây là "bộ não" của hệ thống, nơi các nhà quản trị giám sát hiệu suất thời gian thực thông qua các Dashboard, '
        'Video Wall và Bản đồ số.'
    )
    
    doc.add_heading('2.1 Dashboard Trung tâm', level=2)
    add_image(doc, '02_dashboard_collapsed.png')
    doc.add_paragraph('Dashboard hiển thị các KPI quan trọng: Sản lượng, NRW, Trạng thái trạm và Cảnh báo.')

    doc.add_heading('2.2 Bản đồ GIS', level=2)
    add_image(doc, '04_gis.png')
    doc.add_paragraph('Tích hợp mạng lưới tuyến ống, trạm bơm và sự cố trên nền bản đồ số OSM.')

    # ── Production ─────────────────────────────────────────────────────────
    doc.add_heading('3. Phân hệ Quản lý Sản xuất (Production)', level=1)
    
    doc.add_heading('3.1 Giám sát SCADA', level=2)
    add_image(doc, '07_scada.png')
    doc.add_paragraph('Giám sát lưu lượng, áp lực và mực nước tại các nhà máy nước Hồng Gai, Bãi Cháy, Cẩm Phả.')

    doc.add_heading('3.2 Lập lịch Bơm (Pump Scheduler)', level=2)
    add_image(doc, '11_scheduler.png')
    doc.add_paragraph('Tối ưu hóa thời gian vận hành bơm dựa trên giá điện EVN và đề xuất từ AI.')

    # ── Engineering ────────────────────────────────────────────────────────
    doc.add_heading('4. Phân hệ Kỹ thuật & Mạng lưới (Engineering)', level=1)
    
    doc.add_heading('4.1 Quản lý Sự cố & Lệnh công tác', level=2)
    add_image(doc, '14_incidents.png')
    doc.add_paragraph('Theo dõi vòng đời sự cố từ tiếp nhận đến xử lý tại công trường qua App di động.')

    doc.add_heading('4.2 Quản lý Thất thoát (NRW)', level=2)
    add_image(doc, '13_nrw.png')
    doc.add_paragraph('Phân tích DMA và cân bằng nước để giảm thiểu tỷ lệ thất thoát nước sạch.')

    # ── AI Center ──────────────────────────────────────────────────────────
    doc.add_heading('5. Trung tâm AI (AI Center)', level=1)
    doc.add_paragraph('Sử dụng trí tuệ nhân tạo để hỗ trợ ra quyết định và phân tích dữ liệu chuyên sâu.')
    
    add_image(doc, '22_aiagent.png')
    doc.add_paragraph('AI Agent: Phân tích dữ liệu bằng ngôn ngữ tự nhiên và phát hiện các bất thường vận hành.')

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    print(f"DONE: {OUT}")

if __name__ == "__main__":
    create_doc()
